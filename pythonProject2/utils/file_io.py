"""
文件读写工具模块
处理图像、参数、结果的读写操作
"""

import os
import json
import csv
import yaml
import pickle
import numpy as np
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Union, Tuple
import cv2
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as ReportImage
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging


class DataExporter:
    """数据导出器，支持多种格式导出"""

    def __init__(self, output_dir: str = "results"):
        """
        初始化数据导出器

        Args:
            output_dir: 输出目录
        """
        self.output_dir = output_dir
        self.ensure_directory_exists(output_dir)

    @staticmethod
    def ensure_directory_exists(directory: str):
        """确保目录存在"""
        os.makedirs(directory, exist_ok=True)

    def export_csv(self, data: List[Dict[str, Any]], filename: str = None) -> str:
        """
        导出数据为CSV文件

        Args:
            data: 数据列表，每个元素是一个字典
            filename: 文件名，如果为None则自动生成

        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"results_{timestamp}.csv"

        filepath = os.path.join(self.output_dir, filename)

        try:
            # 转换数据为DataFrame
            df = pd.DataFrame(data)

            # 保存为CSV
            df.to_csv(filepath, index=False, encoding='utf-8-sig')

            logging.info(f"数据已导出到CSV文件: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"导出CSV文件失败: {str(e)}")
            raise

    def export_excel(self, data: List[Dict[str, Any]], filename: str = None) -> str:
        """
        导出数据为Excel文件

        Args:
            data: 数据列表
            filename: 文件名

        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"results_{timestamp}.xlsx"

        filepath = os.path.join(self.output_dir, filename)

        try:
            df = pd.DataFrame(data)

            # 使用ExcelWriter保存多个sheet
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='检测结果', index=False)

                # 添加统计信息sheet
                stats_df = self._calculate_statistics(data)
                stats_df.to_excel(writer, sheet_name='统计信息', index=False)

            logging.info(f"数据已导出到Excel文件: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"导出Excel文件失败: {str(e)}")
            raise

    def export_json(self, data: Dict[str, Any], filename: str = None) -> str:
        """
        导出数据为JSON文件

        Args:
            data: 数据字典
            filename: 文件名

        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"results_{timestamp}.json"

        filepath = os.path.join(self.output_dir, filename)

        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

            logging.info(f"数据已导出到JSON文件: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"导出JSON文件失败: {str(e)}")
            raise

    def export_pickle(self, data: Any, filename: str = None) -> str:
        """
        导出数据为pickle文件

        Args:
            data: 任意Python对象
            filename: 文件名

        Returns:
            导出的文件路径
        """
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"data_{timestamp}.pkl"

        filepath = os.path.join(self.output_dir, filename)

        try:
            with open(filepath, 'wb') as f:
                pickle.dump(data, f)

            logging.info(f"数据已导出到pickle文件: {filepath}")
            return filepath

        except Exception as e:
            logging.error(f"导出pickle文件失败: {str(e)}")
            raise

    def _calculate_statistics(self, data: List[Dict[str, Any]]) -> pd.DataFrame:
        """计算统计信息"""
        if not data:
            return pd.DataFrame()

        df = pd.DataFrame(data)
        stats = []

        # 数值型列的统计
        numeric_cols = df.select_dtypes(include=[np.number]).columns

        for col in numeric_cols:
            stats.append({
                '指标': f'{col}_平均值',
                '值': f"{df[col].mean():.4f}",
                '单位': self._get_unit(col)
            })
            stats.append({
                '指标': f'{col}_标准差',
                '值': f"{df[col].std():.4f}",
                '单位': self._get_unit(col)
            })
            stats.append({
                '指标': f'{col}_最小值',
                '值': f"{df[col].min():.4f}",
                '单位': self._get_unit(col)
            })
            stats.append({
                '指标': f'{col}_最大值',
                '值': f"{df[col].max():.4f}",
                '单位': self._get_unit(col)
            })

        # 合格率统计（如果有is_qualified列）
        if 'is_qualified' in df.columns:
            total = len(df)
            passed = df['is_qualified'].sum()
            failed = total - passed
            pass_rate = (passed / total * 100) if total > 0 else 0

            stats.append({
                '指标': '总检测数',
                '值': str(total),
                '单位': '个'
            })
            stats.append({
                '指标': '合格数',
                '值': str(passed),
                '单位': '个'
            })
            stats.append({
                '指标': '不合格数',
                '值': str(failed),
                '单位': '个'
            })
            stats.append({
                '指标': '合格率',
                '值': f"{pass_rate:.2f}",
                '单位': '%'
            })

        return pd.DataFrame(stats)

    @staticmethod
    def _get_unit(column_name: str) -> str:
        """根据列名获取单位"""
        unit_map = {
            'concentricity': '‰',
            'eccentricity_mm': 'mm',
            'eccentricity_px': 'px',
            'radius': 'px',
            'processing_time': 's',
            'tolerance': 'mm',
            'x': 'px',
            'y': 'px'
        }

        for key, unit in unit_map.items():
            if key in column_name.lower():
                return unit

        return ''


# ====================== 图像文件处理 ======================

def load_image(filepath: str, color_mode: str = 'rgb') -> np.ndarray:
    """
    加载图像文件

    Args:
        filepath: 图像文件路径
        color_mode: 颜色模式，'rgb'或'bgr'

    Returns:
        numpy数组表示的图像
    """
    try:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"图像文件不存在: {filepath}")

        # 使用OpenCV加载图像
        image = cv2.imread(filepath, cv2.IMREAD_UNCHANGED)

        if image is None:
            raise ValueError(f"无法读取图像文件: {filepath}")

        # 转换为RGB格式
        if color_mode.lower() == 'rgb':
            if len(image.shape) == 3 and image.shape[2] == 3:
                image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            elif len(image.shape) == 3 and image.shape[2] == 4:
                image = cv2.cvtColor(image, cv2.COLOR_BGRA2RGBA)

        return image

    except Exception as e:
        logging.error(f"加载图像失败: {str(e)}")
        raise


def save_image(image: np.ndarray, filepath: str,
               quality: int = 95, create_dir: bool = True) -> str:
    """
    保存图像文件

    Args:
        image: numpy数组表示的图像
        filepath: 保存路径
        quality: JPEG质量（1-100）
        create_dir: 是否创建目录

    Returns:
        保存的文件路径
    """
    try:
        # 确保目录存在
        if create_dir:
            directory = os.path.dirname(filepath)
            if directory:
                os.makedirs(directory, exist_ok=True)

        # 确保图像是uint8类型
        if image.dtype != np.uint8:
            if image.max() <= 1.0:
                image = (image * 255).astype(np.uint8)
            else:
                image = image.astype(np.uint8)

        # 根据扩展名确定保存格式
        ext = os.path.splitext(filepath)[1].lower()

        if ext in ['.jpg', '.jpeg']:
            # JPEG格式
            params = [cv2.IMWRITE_JPEG_QUALITY, quality]
            # 如果是RGBA，转换为RGB
            if len(image.shape) == 3 and image.shape[2] == 4:
                image = cv2.cvtColor(image, cv2.COLOR_RGBA2RGB)
            cv2.imwrite(filepath, cv2.cvtColor(image, cv2.COLOR_RGB2BGR), params)

        elif ext == '.png':
            # PNG格式
            params = [cv2.IMWRITE_PNG_COMPRESSION, 9 - int(quality / 100 * 9)]
            if len(image.shape) == 3 and image.shape[2] == 3:
                image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
            elif len(image.shape) == 3 and image.shape[2] == 4:
                image = cv2.cvtColor(image, cv2.COLOR_RGBA2BGRA)
            cv2.imwrite(filepath, image, params)

        else:
            # 其他格式
            if len(image.shape) == 3 and image.shape[2] == 3:
                image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
            cv2.imwrite(filepath, image)

        logging.info(f"图像已保存: {filepath}")
        return filepath

    except Exception as e:
        logging.error(f"保存图像失败: {str(e)}")
        raise


def save_image_with_annotations(image: np.ndarray, filepath: str,
                                annotations: Dict[str, Any] = None,
                                draw_circles: bool = True,
                                draw_text: bool = True) -> str:
    """
    保存带标注的图像

    Args:
        image: 原始图像
        filepath: 保存路径
        annotations: 标注信息
        draw_circles: 是否绘制圆
        draw_text: 是否绘制文本

    Returns:
        保存的文件路径
    """
    try:
        # 创建图像副本
        annotated_image = image.copy()

        if annotations and (draw_circles or draw_text):
            # 绘制标注
            if draw_circles and 'circles' in annotations:
                circles = annotations['circles']
                if isinstance(circles, list) and len(circles) >= 2:
                    # 绘制内圆
                    inner_circle = circles[0]
                    if len(inner_circle) == 3:
                        x, y, r = map(int, inner_circle)
                        cv2.circle(annotated_image, (x, y), r, (0, 255, 0), 2)
                        cv2.circle(annotated_image, (x, y), 3, (0, 255, 0), -1)

                    # 绘制外圆
                    outer_circle = circles[1]
                    if len(outer_circle) == 3:
                        x, y, r = map(int, outer_circle)
                        cv2.circle(annotated_image, (x, y), r, (255, 0, 0), 2)
                        cv2.circle(annotated_image, (x, y), 3, (255, 0, 0), -1)

                    # 绘制中心连线
                    if len(circles[0]) == 3 and len(circles[1]) == 3:
                        x1, y1, _ = map(int, circles[0])
                        x2, y2, _ = map(int, circles[1])
                        cv2.line(annotated_image, (x1, y1), (x2, y2), (255, 255, 0), 2)

            if draw_text and 'results' in annotations:
                results = annotations['results']
                font = cv2.FONT_HERSHEY_SIMPLEX
                y_offset = 30

                for key, value in results.items():
                    if isinstance(value, (int, float)):
                        text = f"{key}: {value:.3f}"
                    else:
                        text = f"{key}: {value}"

                    cv2.putText(annotated_image, text, (10, y_offset),
                                font, 0.7, (255, 255, 255), 2)
                    y_offset += 25

        # 保存图像
        return save_image(annotated_image, filepath)

    except Exception as e:
        logging.error(f"保存标注图像失败: {str(e)}")
        raise


# ====================== 配置文件处理 ======================

def save_parameters_to_yaml(parameters: Dict[str, Any],
                            filepath: str,
                            create_dir: bool = True) -> str:
    """
    保存参数到YAML文件

    Args:
        parameters: 参数字典
        filepath: YAML文件路径
        create_dir: 是否创建目录

    Returns:
        保存的文件路径
    """
    try:
        if create_dir:
            directory = os.path.dirname(filepath)
            if directory:
                os.makedirs(directory, exist_ok=True)

        with open(filepath, 'w', encoding='utf-8') as f:
            yaml.dump(parameters, f, default_flow_style=False, allow_unicode=True)

        logging.info(f"参数已保存到YAML文件: {filepath}")
        return filepath

    except Exception as e:
        logging.error(f"保存YAML文件失败: {str(e)}")
        raise


def load_parameters_from_yaml(filepath: str) -> Dict[str, Any]:
    """
    从YAML文件加载参数

    Args:
        filepath: YAML文件路径

    Returns:
        参数字典
    """
    try:
        if not os.path.exists(filepath):
            logging.warning(f"YAML文件不存在: {filepath}")
            return {}

        with open(filepath, 'r', encoding='utf-8') as f:
            parameters = yaml.safe_load(f)

        return parameters or {}

    except Exception as e:
        logging.error(f"加载YAML文件失败: {str(e)}")
        raise


# ====================== 结果文件处理 ======================

def save_results_to_csv(results: Union[Dict[str, Any], List[Dict[str, Any]]],
                        filepath: str,
                        mode: str = 'w') -> str:
    """
    保存结果到文件（现在默认使用Excel格式）
    为了兼容性，函数名保持不变，但实际生成Excel文件
    """
    try:
        # 如果是单个结果，使用新的Excel函数
        if isinstance(results, dict):
            save_dir = os.path.dirname(filepath) if os.path.dirname(filepath) else "./data/results/"
            return save_concentricity_result_excel(results, save_dir)

        # 批量结果处理（同样使用Excel格式）
        directory = os.path.dirname(filepath)
        if directory:
            os.makedirs(directory, exist_ok=True)

        if isinstance(results, dict):
            results_list = [results]
        else:
            results_list = results

        # 批量数据也保存为Excel
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 如果用户指定的路径是.csv，改为.xlsx
        if filepath.endswith('.csv'):
            filepath = filepath[:-4] + '.xlsx'
        elif not filepath.endswith('.xlsx'):
            filepath = filepath + '.xlsx'

        # 创建DataFrame
        df = pd.DataFrame(results_list)

        # 保存为Excel
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='检测数据', index=False)

            # 添加统计信息sheet
            if len(results_list) > 0:
                stats_data = {
                    '统计项': ['总检测数', '合格数', '不合格数', '合格率'],
                    '数值': [
                        len(results_list),
                        sum(1 for d in results_list if d.get('is_qualified', False)),
                        sum(1 for d in results_list if not d.get('is_qualified', False)),
                        f"{sum(1 for d in results_list if d.get('is_qualified', False)) / len(results_list) * 100:.2f}%"
                    ]
                }
                stats_df = pd.DataFrame(stats_data)
                stats_df.to_excel(writer, sheet_name='统计信息', index=False)

        logging.info(f"批量结果已保存到Excel文件: {filepath}")
        return filepath

    except Exception as e:
        logging.error(f"保存文件失败: {str(e)}")
        raise


def save_results_to_excel(results: Union[Dict[str, Any], List[Dict[str, Any]]],
                          filepath: str) -> str:
    """
    保存结果到Excel文件

    Args:
        results: 单个结果字典或结果列表
        filepath: Excel文件路径

    Returns:
        保存的文件路径
    """
    try:
        # 使用DataExporter
        exporter = DataExporter(os.path.dirname(filepath))

        if isinstance(results, dict):
            results_list = [results]
        else:
            results_list = results

        filename = os.path.basename(filepath)
        return exporter.export_excel(results_list, filename)

    except Exception as e:
        logging.error(f"保存Excel文件失败: {str(e)}")
        raise


# ====================== 报告生成 ======================

def export_report_to_pdf(data: Dict[str, Any], filepath: str,
                         title: str = "同心度检测报告",
                         company: str = "西南科技大学信息与控制工程学院",
                         author: str = "机械零件同心度检测系统") -> str:
    """
    导出报告为PDF文件（修复版）
    关键改进：确保路径正确，自动创建目录，提供清晰反馈
    """
    try:
        # 确保目录存在
        directory = os.path.dirname(filepath)
        if directory:
            os.makedirs(directory, exist_ok=True)
            logging.info(f"确保PDF保存目录存在: {directory}")

        # 打印实际保存路径（便于调试）
        print(f"[DEBUG] PDF将保存到: {os.path.abspath(filepath)}")

        # 创建PDF文档
        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        styles = getSampleStyleSheet()
        story = []

        # 标题
        title_style = styles['Heading1']
        title_style.alignment = 1  # 居中
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # 基本信息
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Paragraph(f"生成单位: {company}", styles['Normal']))
        story.append(Paragraph(f"生成系统: {author}", styles['Normal']))
        story.append(Spacer(1, 20))

        # 检测结果汇总
        story.append(Paragraph("检测结果汇总", styles['Heading2']))

        # 从数据中提取关键信息
        if 'results' in data:
            results = data['results']
        else:
            # 尝试从原数据中提取
            results = {
                '检测时间': data.get('detection_time', 'N/A'),
                '图像尺寸': data.get('image_size', 'N/A'),
                '实际偏差': f"{data.get('error_mm', 0):.3f} mm",
                '像素偏差': f"{data.get('pixel_error', 0):.2f} px",
                '相对偏差': f"{data.get('relative_error_percent', 0):.3f}%",
                '合格状态': '合格' if data.get('is_qualified', False) else '不合格'
            }

        # 创建结果表格
        results_data = [['参数', '数值', '单位']]

        # 添加关键检测结果
        key_mapping = {
            '检测时间': ('', ''),
            '图像尺寸': ('', ''),
            '实际偏差': ('mm', '实际偏心距'),
            '像素偏差': ('px', '像素偏心距'),
            '相对偏差': ('%', '相对偏差'),
            '同心度': ('‰', '同心度'),
            '偏心距': ('mm', '偏心距')
        }

        for key, value in results.items():
            if isinstance(value, (int, float)):
                value_str = f"{value:.4f}"
            else:
                value_str = str(value)

            # 确定单位和显示名称
            display_name = key
            unit = ''

            for map_key, (map_unit, map_name) in key_mapping.items():
                if map_key in key:
                    unit = map_unit
                    if map_name:
                        display_name = map_name
                    break

            results_data.append([display_name, value_str, unit])

        # 添加从原始数据中提取的信息
        if 'inner_center' in data:
            inner = data['inner_center']
            if isinstance(inner, dict):
                results_data.append(['螺纹杆中心X', f"{inner.get('x', 0):.1f}", 'px'])
                results_data.append(['螺纹杆中心Y', f"{inner.get('y', 0):.1f}", 'px'])

        if 'outer_circle' in data:
            outer = data['outer_circle']
            if isinstance(outer, dict):
                results_data.append(['外圆中心X', f"{outer.get('x', 0):.1f}", 'px'])
                results_data.append(['外圆中心Y', f"{outer.get('y', 0):.1f}", 'px'])
                results_data.append(['外圆半径', f"{outer.get('radius', 0):.1f}", 'px'])

        # 创建表格
        table = Table(results_data, colWidths=[2 * inch, 2.5 * inch, 0.8 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F81BD')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#DCE6F1')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#E6E6E6')])
        ]))

        story.append(table)
        story.append(Spacer(1, 20))

        # 结论部分
        story.append(Paragraph("检测结论", styles['Heading2']))
        conclusion = "产品合格，符合质量标准。" if data.get('is_qualified', False) else "产品不合格，超出公差范围。"
        story.append(Paragraph(conclusion, styles['Normal']))

        # 生成PDF
        doc.build(story)

        logging.info(f"PDF报告已成功生成: {filepath}")
        print(f"[SUCCESS] PDF报告已保存到: {os.path.abspath(filepath)}")
        return filepath

    except Exception as e:
        logging.error(f"导出PDF报告失败: {str(e)}")
        print(f"[ERROR] 生成PDF失败: {str(e)}")
        raise


def export_report_to_word(data: Dict[str, Any], filepath: str,
                          title: str = "同心度检测报告") -> str:
    """
    导出报告为Word文件

    Args:
        data: 报告数据
        filepath: Word文件路径
        title: 报告标题

    Returns:
        导出的文件路径
    """
    try:
        # 确保目录存在
        directory = os.path.dirname(filepath)
        if directory:
            os.makedirs(directory, exist_ok=True)

        # 创建Word文档
        doc = Document()

        # 标题
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # 空行

        # 基本信息
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"生成系统: 同心度检测系统 v1.0")

        doc.add_paragraph()  # 空行

        # 检测结果
        if 'results' in data:
            doc.add_paragraph("检测结果").bold = True

            # 创建表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid'

            # 表头
            header_cells = table.rows[0].cells
            header_cells[0].text = "参数"
            header_cells[1].text = "数值"
            header_cells[2].text = "单位"

            # 数据行
            for key, value in data['results'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = key

                if isinstance(value, (int, float)):
                    row_cells[1].text = f"{value:.4f}"
                else:
                    row_cells[1].text = str(value)

                # 确定单位
                unit = ''
                if 'concentricity' in key.lower():
                    unit = '‰'
                elif 'eccentricity' in key.lower():
                    unit = 'mm'
                elif 'radius' in key.lower():
                    unit = 'mm'
                elif 'time' in key.lower():
                    unit = 's'

                row_cells[2].text = unit

        doc.add_paragraph()  # 空行

        # 备注
        if 'remarks' in data:
            doc.add_paragraph("备注").bold = True
            doc.add_paragraph(data['remarks'])

        # 保存文档
        doc.save(filepath)

        logging.info(f"报告已导出为Word: {filepath}")
        return filepath

    except Exception as e:
        logging.error(f"导出Word报告失败: {str(e)}")
        raise


# ====================== 其他工具函数 ======================

def create_directory(directory: str) -> str:
    """
    创建目录（如果不存在）

    Args:
        directory: 目录路径

    Returns:
        创建的目录路径
    """
    try:
        os.makedirs(directory, exist_ok=True)
        return directory
    except Exception as e:
        logging.error(f"创建目录失败: {str(e)}")
        raise


def get_unique_filename(base_name: str, extension: str,
                        directory: str = ".") -> str:
    """
    获取唯一的文件名

    Args:
        base_name: 基础文件名
        extension: 文件扩展名（带点）
        directory: 目录路径

    Returns:
        唯一的文件路径
    """
    try:
        create_directory(directory)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 尝试带时间戳的文件名
        filename = f"{base_name}_{timestamp}{extension}"
        filepath = os.path.join(directory, filename)

        if not os.path.exists(filepath):
            return filepath

        # 如果存在，添加序号
        counter = 1
        while True:
            filename = f"{base_name}_{timestamp}_{counter:03d}{extension}"
            filepath = os.path.join(directory, filename)

            if not os.path.exists(filepath):
                return filepath

            counter += 1

    except Exception as e:
        logging.error(f"生成唯一文件名失败: {str(e)}")
        raise


# ====================== 新增：同心度检测结果保存函数 ======================
def save_concentricity_result_excel(result: Dict[str, Any], save_dir: str = "./data/results/") -> str:
    """
    保存同心度检测结果为Excel文件（核心数据，无格式，适合数据分析）
    仅保存核心数据，不包含处理参数等冗余信息

    Args:
        result: 检测结果字典
        save_dir: 保存目录（默认：./data/results/）

    Returns:
        保存的文件路径
    """
    try:
        # 确保目录存在
        os.makedirs(save_dir, exist_ok=True)

        # 只提取核心数据字段（用于批量分析和处理）
        core_data = {}

        # 1. 检测基本信息
        core_data["检测时间"] = result.get("detection_time", "")
        core_data["图像尺寸"] = result.get("image_size", "")
        core_data["是否合格"] = "是" if result.get("is_qualified", False) else "否"
        core_data["实际偏差(mm)"] = round(result.get("error_mm", 0), 4)
        core_data["像素偏差(px)"] = round(result.get("pixel_error", 0), 2)
        core_data["同心度(‰)"] = round(result.get("concentricity", 0), 4)

        # 2. 内圆坐标
        inner_center = result.get("inner_center", {})
        if isinstance(inner_center, dict):
            core_data["内圆X(px)"] = round(inner_center.get("x", 0), 2)
            core_data["内圆Y(px)"] = round(inner_center.get("y", 0), 2)
        else:
            core_data["内圆X(px)"] = 0.0
            core_data["内圆Y(px)"] = 0.0

        # 3. 外圆坐标
        outer_circle = result.get("outer_circle", {})
        if isinstance(outer_circle, dict):
            core_data["外圆X(px)"] = round(outer_circle.get("x", 0), 2)
            core_data["外圆Y(px)"] = round(outer_circle.get("y", 0), 2)
            core_data["外圆半径(px)"] = round(outer_circle.get("radius", 0), 2)
        else:
            core_data["外圆X(px)"] = 0.0
            core_data["外圆Y(px)"] = 0.0
            core_data["外圆半径(px)"] = 0.0

        # 4. 其他关键参数
        core_data["处理时间(s)"] = round(result.get("processing_time", 0), 3)
        core_data["检测置信度"] = round(result.get("confidence", 0), 3)

        # 生成文件名
        timestamp = result.get("detection_time", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        if isinstance(timestamp, str):
            # 清理时间字符串，使其适合文件名
            timestamp_clean = timestamp.replace(":", "").replace(" ", "_").replace("-", "")
        else:
            timestamp_clean = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 保存为Excel文件
        excel_filename = f"同心度检测结果_{timestamp_clean}.xlsx"
        excel_filepath = os.path.join(save_dir, excel_filename)

        # 创建DataFrame
        df = pd.DataFrame([core_data])

        # 使用openpyxl引擎保存为Excel
        try:
            with pd.ExcelWriter(excel_filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='检测结果', index=False)

                # 获取worksheet对象以调整列宽
                workbook = writer.book
                worksheet = writer.sheets['检测结果']

                # 自动调整所有列宽
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter  # 获取列字母

                    # 计算每列的最大长度
                    for cell in column:
                        try:
                            cell_value = str(cell.value) if cell.value is not None else ""
                            cell_length = len(cell_value)
                            if cell_length > max_length:
                                max_length = cell_length
                        except:
                            pass

                    # 设置列宽
                    adjusted_width = min(max_length + 2, 30)  # 最大宽度30
                    worksheet.column_dimensions[column_letter].width = adjusted_width

            logging.info(f"同心度检测结果已保存为Excel（核心数据）: {excel_filepath}")
            print(f"[INFO] 核心数据Excel文件已保存到: {excel_filepath}")

            return excel_filepath

        except Exception as excel_error:
            logging.error(f"保存Excel文件失败: {str(excel_error)}")
            raise

    except Exception as e:
        logging.error(f"保存同心度结果失败: {str(e)}")
        # 降级处理：保存原始数据
        try:
            fallback_path = os.path.join(save_dir, f"同心度原始数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
            with open(fallback_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
            logging.info(f"原始数据已保存为JSON: {fallback_path}")
            return fallback_path
        except:
            raise Exception(f"无法保存结果: {str(e)}")


# ====================== 新增：快速生成报告函数 ======================
def quick_generate_report(result_data: Dict[str, Any],
                          report_dir: str = "./data/reports/") -> str:
    """
    快速生成同心度检测报告（简化版，自动保存到正确目录）

    Args:
        result_data: 检测结果数据
        report_dir: 报告保存目录（默认：./data/reports/）

    Returns:
        生成的PDF文件路径
    """
    try:
        # 确保报告目录存在
        os.makedirs(report_dir, exist_ok=True)

        # 生成文件名
        timestamp = result_data.get("detection_time", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        if isinstance(timestamp, str):
            timestamp_clean = timestamp.replace(":", "").replace(" ", "_").replace("-", "")
        else:
            timestamp_clean = datetime.now().strftime("%Y%m%d_%H%M%S")

        filename = f"同心度报告_{timestamp_clean}.pdf"
        filepath = os.path.join(report_dir, filename)

        # 生成报告
        return export_report_to_pdf(result_data, filepath)

    except Exception as e:
        logging.error(f"快速生成报告失败: {str(e)}")
        raise


def save_multiple_results_to_excel(results: List[Dict[str, Any]],
                                   filepath: str = None) -> str:
    """
    保存多条检测结果到Excel文件

    Args:
        results: 检测结果列表
        filepath: 保存路径（如果为None则自动生成）

    Returns:
        保存的文件路径
    """
    try:
        if not results:
            raise ValueError("没有结果可保存")

        # 如果未指定文件路径，自动生成
        if filepath is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            save_dir = os.path.join(os.getcwd(), "data", "results")
            os.makedirs(save_dir, exist_ok=True)
            filepath = os.path.join(save_dir, f"批量检测结果_{timestamp}.xlsx")

        # 确保目录存在
        directory = os.path.dirname(filepath)
        if directory:
            os.makedirs(directory, exist_ok=True)

        # 提取核心数据
        core_data_list = []

        for i, result in enumerate(results, 1):
            core_data = {}

            # 1. 检测基本信息
            core_data["序号"] = i
            core_data["检测时间"] = result.get("detection_time", "")
            core_data["是否合格"] = "是" if result.get("is_qualified", False) else "否"
            core_data["实际偏差(mm)"] = round(result.get("error_mm", 0), 4)
            core_data["像素偏差(px)"] = round(result.get("pixel_error", 0), 2)
            core_data["同心度(‰)"] = round(result.get("concentricity", 0), 4)

            # 2. 内圆坐标
            inner_center = result.get("inner_center", {})
            if isinstance(inner_center, dict):
                core_data["内圆X(px)"] = round(inner_center.get("x", 0), 2)
                core_data["内圆Y(px)"] = round(inner_center.get("y", 0), 2)
            else:
                core_data["内圆X(px)"] = 0.0
                core_data["内圆Y(px)"] = 0.0

            # 3. 外圆坐标
            outer_circle = result.get("outer_circle", {})
            if isinstance(outer_circle, dict):
                core_data["外圆X(px)"] = round(outer_circle.get("x", 0), 2)
                core_data["外圆Y(px)"] = round(outer_circle.get("y", 0), 2)
                core_data["外圆半径(px)"] = round(outer_circle.get("radius", 0), 2)
            else:
                core_data["外圆X(px)"] = 0.0
                core_data["外圆Y(px)"] = 0.0
                core_data["外圆半径(px)"] = 0.0

            core_data_list.append(core_data)

        # 创建DataFrame
        df = pd.DataFrame(core_data_list)

        # 保存为Excel
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='批量检测结果', index=False)

            # 添加统计信息sheet
            total_count = len(results)
            qualified_count = sum(1 for r in results if r.get('is_qualified', False))
            qualified_rate = (qualified_count / total_count * 100) if total_count > 0 else 0

            stats_data = {
                '统计项': ['总检测数', '合格数', '不合格数', '合格率'],
                '数值': [
                    total_count,
                    qualified_count,
                    total_count - qualified_count,
                    f"{qualified_rate:.2f}%"
                ]
            }
            stats_df = pd.DataFrame(stats_data)
            stats_df.to_excel(writer, sheet_name='统计信息', index=False)

            # 调整列宽
            workbook = writer.book
            worksheet = writer.sheets['批量检测结果']

            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter

                for cell in column:
                    try:
                        cell_value = str(cell.value) if cell.value is not None else ""
                        cell_length = len(cell_value)
                        if cell_length > max_length:
                            max_length = cell_length
                    except:
                        pass

                adjusted_width = min(max_length + 2, 30)
                worksheet.column_dimensions[column_letter].width = adjusted_width

        logging.info(f"批量结果已保存到Excel文件: {filepath}")
        return filepath

    except Exception as e:
        logging.error(f"保存批量结果失败: {str(e)}")
        raise